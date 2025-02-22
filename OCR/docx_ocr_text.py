from docx import Document
import pytesseract
from PIL import Image
import io
import os
import tempfile
import shutil
import re

# Path to Tesseract OCR executable
#pytesseract.pytesseract.tesseract_cmd = r'C:\\Program Files\\Tesseract-OCR\\tesseract.exe'


# Function to extract images from a Word document
def extract_images_from_docx(docx_path, output_dir):
    doc = Document(docx_path)
    
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    image_paths = []
    for i, rel in enumerate(doc.part.rels.values()):
        if "image" in rel.target_ref:
            image_data = rel.target_part.blob
            image_path = os.path.join(output_dir, f'image_{i + 1}.png')
            with open(image_path, "wb") as f:
                f.write(image_data)
            image_paths.append(image_path)
    return image_paths
# Function to perform OCR on images
def perform_ocr_on_images(image_paths):
    ocr_results = {}
    for image_path in image_paths:
        try:
            image = Image.open(image_path)
            text = pytesseract.image_to_string(image)
            ocr_results[image_path] = text
        except Exception as e:
            print(f"Error processing {image_path}: {e}")

    return ocr_results

# Function to extract text from word file
def ocr_from_word_file(docx_path, output_dir):
    pytesseract.pytesseract.tesseract_cmd = r'C:\\Program Files\\Tesseract-OCR\\tesseract.exe'
    print("Extracting images...")
    
    #image_paths = extract_images_from_docx(docx_path, output_dir)
    doc = Document(docx_path)
    ocr_results=""
    if doc.paragraphs: 
        print("in here")
        fullText = []
        for para in doc.paragraphs:
            print("printing para")
            #print(para.text)
            fullText.append(para.text)
        ocr_results += '\n'.join(fullText)
    '''
    if image_paths:
        print("Performing OCR...")
        ocr_results += perform_ocr_on_images(image_paths)
    '''
    return ocr_results

output_directory = "extracted_images"  # Directory to save extracted images
def ocr_docx(filepath):
    global output_directory
    ocr_texts = ocr_from_word_file(filepath, output_directory)
    '''
    with open("docx_text","w") as file:
        file.write(ocr_texts)
    shutil.rmtree(output_directory)
    '''
    #print(ocr_texts)
    cleaned_text = re.sub(r'[^A-Za-z0-9\s.]', '', ocr_texts)
    text = re.sub(r'[^\w\s.,!?]', '', cleaned_text)
    cleaned_text1 = re.sub(r'[^\x20-\x7E\n]', '', text)
    return cleaned_text1
    # Print the OCR results

'''
word_ocr=ocr_docx("C:\\Users\\Deepak J Bhat\\Downloads\\newdoc.docx")
with open("ocr_text_word","w",encoding="utf-8") as file:
    file .write(word_ocr)
'''

